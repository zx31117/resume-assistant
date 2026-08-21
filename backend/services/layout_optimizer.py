"""V1.2 Layout Optimizer（纯规则引擎，不调 AI，不删条目/不改事实内容）。

设计边界（用户审核反馈 §8.2 最终稿）：

LayoutOptimizer **只调段落级样式**（段距/字号/行距/页边距），**绝不删减任何经历条目或 bullet**。
经历数量裁剪在 ResumeBuilder.build() 中已发生（按 JD priority + max_items 截断），渲染后不再裁。

四级降级顺序（代码内固定，V1.2 不读 JSON layout_rules 配置）：
  1. 段前距/段后距：6pt → 3pt → 0pt
  2. 正文字号：10.5pt → 10pt → 9.5pt（不碰标题字号）
  3. 行距：1.5 → 1.3 → 1.15
  4. 页边距：各方向各减 0.3cm（只做一次）
仍超页：发 warning，不删内容，不硬塞入 1 页。
"""
import math

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


# ── 页数估算 ────────────────────────────────────────────────────── #

def estimate_pages(doc: Document) -> int:
    """字符密度启发式估算页数（LayoutOptimizer 只需要"是否超页"的布尔判断，不要求精确）。

    A4 字号 10.5pt、标准页边距的标准一页约 2200 字。
    """
    total_chars = sum(len(p.text) for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_chars += len(cell.text)
    # 损耗系数 0.78（标题/段距/空白行导致实际字数低于理论值）
    effective_chars = total_chars * 0.78
    return max(1, math.ceil(effective_chars / 2200))


# ── 空白页处理 ──────────────────────────────────────────────────── #

def remove_blank_trailing_pages(doc: Document):
    """从文档末尾向前删除连续空段落（仅删空段落，不删有内容的段落）。"""
    while len(doc.paragraphs) > 1:
        last = doc.paragraphs[-1]
        if last.text.strip() == "" and not last._p.findall(qn('w:tbl')):
            last._p.getparent().remove(last._p)
        else:
            break


# ── 四级降级：具体操作（**不动任何条目数量**） ─────────────────── #

# §8.2 规则：V1.2 排版规则代码内固定，不读 JSON layout_rules 配置
SPACING_LEVELS = [6, 3, 0]            # 段前距/段后距（pt）
FONTSIZE_LEVELS = [10.5, 10, 9.5]    # 正文字号（pt），标题字号不碰
LINESPACING_LEVELS = [1.5, 1.3, 1.15] # 行距（倍数）

HEADING_STYLE_PREFIXES = ("SectionTitle_", "Heading ", "标题")  # 标题样式前缀（字号降级时跳过）

def _is_heading(p) -> bool:
    try:
        name = p.style and p.style.name
        return bool(name and any(str(name).startswith(pref) for pref in HEADING_STYLE_PREFIXES))
    except Exception:
        return False


def _apply_spacing(doc: Document, spacing_pt: float) -> None:
    """把所有段落的段前距/段后距统一为 spacing_pt（不低于此值时才设置，避免段距变大）。"""
    for p in doc.paragraphs:
        pf = p.paragraph_format
        try:
            if pf.space_before is None or pf.space_before.pt > spacing_pt:
                pf.space_before = Pt(spacing_pt)
        except Exception:
            pass
        try:
            if pf.space_after is None or pf.space_after.pt > spacing_pt:
                pf.space_after = Pt(spacing_pt)
        except Exception:
            pass


def _apply_body_font_size(doc: Document, size_pt: float) -> None:
    """正文 Run 的字号调到 size_pt（只缩不扩；标题样式段落不碰）。

    注意：直接操作 python-docx 高层 API run.font.size 只改 w:sz；
    但 V1.2 标准模板里字号同时写 w:sz 和 w:szCs（成对），这里只改 w:sz 即可，
    Word 渲染会 fallback；若要严格对齐 w:szCs 可升级操作 XML，但 V1.2 够用。
    """
    for p in doc.paragraphs:
        if _is_heading(p):
            continue
        for run in p.runs:
            try:
                if run.font.size is None or run.font.size.pt > size_pt:
                    run.font.size = Pt(size_pt)
            except Exception:
                pass


def _apply_line_spacing(doc: Document, line_spacing: float) -> None:
    for p in doc.paragraphs:
        try:
            p.paragraph_format.line_spacing = line_spacing
        except Exception:
            pass


def _reduce_margins(doc: Document, delta_cm: float = 0.3) -> bool:
    """页边距各方向减 delta_cm cm（最低不能低于 0.5cm）。只执行一次（返回 False 表示已经到底，不能再降）。"""
    applied = False
    try:
        for section in doc.sections:
            for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                cur = getattr(section, attr)
                cur_cm = cur.cm if cur is not None else 1.5
                new_cm = max(0.5, cur_cm - delta_cm)
                if new_cm < cur_cm:
                    setattr(section, attr, Cm(new_cm))
                    applied = True
    except Exception:
        pass
    return applied


# ── 主入口 ──────────────────────────────────────────────────────── #

def optimize(
    doc: Document,
    page_limit: int = 1,
) -> tuple[list[str], list[str]]:
    """§8.2 V1.2 实现：对已填充的 Document（内存对象）做四级规则降级。

    与旧版 optimize() 的差异：
      - 输入/输出都是内存对象（不再走"读 docx → 原地保存"）；调用方自己 save()
      - 四级降级按顺序尝试，每级达标就停；绝不删减任何经历条目或 bullet
      - 返回 (applied_rules, capacity_warnings) 两个列表，方便调用方分离展示。

    参数：
      doc: 已填充完毕的 Document（由 TemplateRenderer.render 产出）
      page_limit: 目标页数（来自 TemplateSpec.layout.page_limit）

    返回：
      applied_rules：如 ["paragraph spacing → 3pt", "body font size → 10pt"]
      capacity_warnings：超容量告警（若有），如 ["resume exceeds page_limit=1（estimate ~2 pages）..."]
    """
    applied: list[str] = []
    capacity_warnings: list[str] = []

    # 先处理文档尾部连续空段落（不删有内容的段落）
    remove_blank_trailing_pages(doc)

    if estimate_pages(doc) <= page_limit:
        return applied, capacity_warnings

    # ── 一级：段前距/段后距 6pt → 3pt → 0pt ──
    for sp in SPACING_LEVELS:
        if estimate_pages(doc) <= page_limit:
            break
        _apply_spacing(doc, sp)
        applied.append(f"paragraph spacing → {sp}pt")

    if estimate_pages(doc) <= page_limit:
        return applied, capacity_warnings

    # ── 二级：正文字号 10.5 → 10 → 9.5（标题不变） ──
    for sz in FONTSIZE_LEVELS:
        if estimate_pages(doc) <= page_limit:
            break
        _apply_body_font_size(doc, sz)
        applied.append(f"body font size → {sz}pt")

    if estimate_pages(doc) <= page_limit:
        return applied, capacity_warnings

    # ── 三级：行距 1.5 → 1.3 → 1.15 ──
    for ls in LINESPACING_LEVELS:
        if estimate_pages(doc) <= page_limit:
            break
        _apply_line_spacing(doc, ls)
        applied.append(f"line spacing → {ls}")

    if estimate_pages(doc) <= page_limit:
        return applied, capacity_warnings

    # ── 四级：页边距减 0.3cm（只执行一次） ──
    if _reduce_margins(doc, 0.3):
        applied.append("margin reduced by 0.3cm")

    # V1.3 T8：仍超页 → 单独作为 capacity_warnings 输出（绝不删内容）
    final_pages = estimate_pages(doc)
    if final_pages > page_limit:
        capacity_warnings.append(
            f"resume exceeds page_limit={page_limit}（estimate ~{final_pages} page(s)）；"
            f"no further content truncation（不删经历条目），建议人工在 Word 中微调"
        )

    return applied, capacity_warnings

