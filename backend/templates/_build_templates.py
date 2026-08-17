"""构建 pm_template.docx（V1.2 标准化模板 — 参考 input/用户上传PDF.pdf 布局）。

【设计依据】—— 对照 用户上传PDF.pdf 提取结果：
  页面：A4，页边距 T=0.92cm B=1.39cm L=1.13cm R=1.09cm
  字体：Microsoft YaHei（微软雅黑），粗体由 bold 属性控制
  字号：
    姓名 20pt 加粗 #0D0D0D  |  章节标题 12pt 加粗 #0D0D0D（带底分隔线）
    经历标题行 10.6pt 加粗 #262626  |  正文 bullets 10.6pt 常规 #262626
    联系方式/求职意向 10pt 常规 #262626
  章节顺序（与 PDF 完全一致）：
    1. Profile（姓名 + 求职意向 + 联系方式 + 右上角照片占位框）— 全部左对齐
    2. 教育背景（Education）
    3. 实习经历（Work）
    4. 项目经历（Project）
    5. 技能专长（Skills）
    6. 荣誉奖项（Awards，可选）
    7. 自我评价（Summary，可选）
  关键排版：
    * 经历标题行三列 Tab 布局：时间(左) \t 学校/公司(居中) \t 专业/职位(右)
    * bullet 用 ⚫（U+26AB）前缀，不是 •
    * 章节标题下方底边框 = 分隔线
    * 照片框：右上角浮动 anchored drawing（2.375×2.9cm，距上0.58cm 距右1.63cm）
    * 无关键词加粗（完全按 PDF，bullet 行全部常规字体）
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "pm_template.docx")

# ── PDF 实测参数 ──────────────────────────────────────────────────
FONT_CN = "微软雅黑"
FONT_EN = "Microsoft YaHei"
COLOR_TITLE = "0D0D0D"   # 姓名/章节标题（接近纯黑）
COLOR_BODY = "262626"   # 经历标题行/正文（深灰）

# 页边距（cm）
MARGIN_TOP = 0.92
MARGIN_BOTTOM = 1.39
MARGIN_LEFT = 1.13
MARGIN_RIGHT = 1.09
# 文字区宽度 = 21 - 1.13 - 1.09 = 18.78cm；中点 = 9.39cm
TEXT_WIDTH_CM = 21.0 - MARGIN_LEFT - MARGIN_RIGHT
TAB_CENTER_CM = TEXT_WIDTH_CM / 2          # 9.39
TAB_RIGHT_CM = TEXT_WIDTH_CM               # 18.78

# 字号
SZ_NAME = 20.0
SZ_TITLE = 12.0
SZ_BODY = 10.6
SZ_CONTACT = 10.0

# 照片框（PDF 实测）
PHOTO_W_CM = 2.375
PHOTO_H_CM = 2.9
PHOTO_TOP_CM = 0.58                       # 距页面顶部
# 照片框页面绝对定位（relativeFrom=page），距右1.63cm = 距页面右边缘
PHOTO_FROM_LEFT_CM = 21.0 - 1.63 - PHOTO_W_CM  # ≈16.995

# 行距（PDF 实测：正文行高约 18pt）
LINE_BODY = 18.0


def set_run_font(run, *, cn_font=FONT_CN, en_font=FONT_EN,
                 size_pt=None, bold=None, color=None, hint="eastAsia"):
    """给 Run 设置中英字体 + 字号 + 加粗 + 颜色。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    rFonts.set(qn('w:cs'), en_font)
    if hint:
        rFonts.set(qn('w:hint'), hint)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = OxmlElement('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
    if bold is not None:
        run.font.bold = bold
        b = rPr.find(qn('w:b'))
        if bold:
            if b is None:
                b = OxmlElement('w:b')
                rPr.append(b)
            b.set(qn('w:val'), '1')
        else:
            if b is not None:
                rPr.remove(b)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def ensure_style(doc, name, *, based_on="Normal",
                 font_cn=FONT_CN, font_en=FONT_EN,
                 size_pt=SZ_BODY, bold=False, color=None,
                 space_before=0, space_after=3,
                 first_line_indent_chars=0, alignment=None,
                 line_pt=None):
    """创建/获取命名样式。"""
    try:
        style = doc.styles[name]
    except KeyError:
        style = doc.styles.add_style(name, 1)
        style.base_style = doc.styles[based_on] if based_on in doc.styles else doc.styles["Normal"]
    pf = style.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_line_indent_chars:
        pf.first_line_indent = Pt(size_pt * first_line_indent_chars)
    if alignment is not None:
        pf.alignment = alignment
    if line_pt is not None:
        pf.line_spacing = Pt(line_pt)
        pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    f = style.font
    f.name = font_en
    f.size = Pt(size_pt)
    f.bold = bold
    if color is not None:
        f.color.rgb = RGBColor.from_string(color)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_cn)
    rFonts.set(qn('w:ascii'), font_en)
    rFonts.set(qn('w:hAnsi'), font_en)
    rFonts.set(qn('w:cs'), font_en)
    rFonts.set(qn('w:hint'), "eastAsia")
    return style


def _run(p, text, *, size_pt=SZ_BODY, bold=False, color=None):
    r = p.add_run(text)
    set_run_font(r, size_pt=size_pt, bold=bold, color=color)
    return r


def _ph(p, placeholder, *, size_pt=SZ_BODY, bold=False, color=None):
    """占位符独占 Run（R5）。"""
    r = p.add_run(placeholder)
    set_run_font(r, size_pt=size_pt, bold=bold, color=color)
    return r


def _add_bottom_border(p, sz=6, color=COLOR_TITLE, space=2):
    """给段落加底边框（章节标题分隔线）。sz 单位 1/8pt，6=0.75pt。"""
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), str(space))
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)


def _add_tab_stops(p):
    """经历标题行三列 Tab：中点居中 + 右边右对齐。"""
    ts = p.paragraph_format.tab_stops
    ts.add_tab_stop(Cm(TAB_CENTER_CM), WD_TAB_ALIGNMENT.CENTER)
    ts.add_tab_stop(Cm(TAB_RIGHT_CM), WD_TAB_ALIGNMENT.RIGHT)


def _add_photo_placeholder(doc):
    """右上角浮动照片占位框（anchored drawing，浅灰填充 + 细边框）。"""
    w_emu = int(Cm(PHOTO_W_CM))
    h_emu = int(Cm(PHOTO_H_CM))
    x_emu = int(Cm(PHOTO_FROM_LEFT_CM))
    y_emu = int(Cm(PHOTO_TOP_CM))
    # 取文档第一个段落插入 drawing run（照片浮动，不受文字流影响）
    first_p = doc.paragraphs[0]
    xml = (
        f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        f'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        f'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<w:drawing>'
        f'<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" '
        f'relativeHeight="1" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">'
        f'<wp:simplePos x="0" y="0"/>'
        f'<wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>'
        f'<wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>'
        f'<wp:extent cx="{w_emu}" cy="{h_emu}"/>'
        f'<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        f'<wp:wrapNone/>'
        f'<wp:docPr id="100" name="PhotoPlaceholder"/>'
        f'<wp:cNvGraphicFramePr/>'
        f'<a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f'<wps:wsp><wps:cNvSpPr txBox="0"/>'
        f'<wps:spPr>'
        f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>'
        f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
        f'<a:solidFill><a:srgbClr val="F2F2F2"/></a:solidFill>'
        f'<a:ln w="6350"><a:solidFill><a:srgbClr val="BFBFBF"/></a:solidFill></a:ln>'
        f'</wps:spPr>'
        f'<wps:bodyPr rot="0" spcFirstLastPara="0" vert="horz" wrap="square" '
        f'lIns="0" tIns="0" rIns="0" bIns="0" anchor="ctr"/>'
        f'</wps:wsp></a:graphicData></a:graphic>'
        f'</wp:anchor></w:drawing></w:r>'
    )
    drawing_el = parse_xml(xml)
    first_p._p.append(drawing_el)


def _make_item_title(doc, style_name, cells):
    """经历标题行：时间 \\t 学校/公司 \\t 专业/职位（三列 Tab）。
    cells = [(ph_key, suffix), ...]  suffix 含分隔符如 '-' '\\t' '（' '）'
    """
    p = doc.add_paragraph(style=style_name)
    _add_tab_stops(p)
    for ph_key, suffix in cells:
        _ph(p, "{{" + ph_key + "}}", size_pt=SZ_BODY, bold=True, color=COLOR_BODY)
        if suffix:
            _run(p, suffix, size_pt=SZ_BODY, bold=True, color=COLOR_BODY)
    return p


def _make_bullet(doc, style_name, ph_key):
    """bullet 段落：⚫ + 占位符（⚫ 常规，占位符常规）。"""
    p = doc.add_paragraph(style=style_name)
    _run(p, "⚫", size_pt=SZ_BODY, bold=False, color=COLOR_BODY)
    _ph(p, "{{" + ph_key + "}}", size_pt=SZ_BODY, bold=False, color=COLOR_BODY)
    return p


def build():
    doc = Document()

    # ── 页面设置（PDF 实测页边距）──
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(MARGIN_TOP)
        section.bottom_margin = Cm(MARGIN_BOTTOM)
        section.left_margin = Cm(MARGIN_LEFT)
        section.right_margin = Cm(MARGIN_RIGHT)

    # ── Normal 默认字号 ──
    ensure_style(doc, "Normal", size_pt=SZ_BODY)

    # ── Profile 系列（全部左对齐，与 PDF 一致）──
    ensure_style(doc, "Profile_Name",   size_pt=SZ_NAME, bold=True,  color=COLOR_TITLE,
                 space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=24)
    ensure_style(doc, "Profile_Target", size_pt=SZ_CONTACT, bold=False, color=COLOR_BODY,
                 space_before=2, space_after=2, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=14)
    ensure_style(doc, "Profile_Line",   size_pt=SZ_CONTACT, bold=False, color=COLOR_BODY,
                 space_before=0, space_after=4, alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=14)

    # ── SectionTitle 系列（12pt 粗，底分隔线在段落级加）──
    TITLE_BEFORE = 8
    TITLE_AFTER = 4
    ensure_style(doc, "SectionTitle_Education", size_pt=SZ_TITLE, bold=True, color=COLOR_TITLE,
                 space_before=TITLE_BEFORE, space_after=TITLE_AFTER,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16)
    ensure_style(doc, "SectionTitle_Work", size_pt=SZ_TITLE, bold=True, color=COLOR_TITLE,
                 space_before=TITLE_BEFORE, space_after=TITLE_AFTER,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16)
    ensure_style(doc, "SectionTitle_Project", size_pt=SZ_TITLE, bold=True, color=COLOR_TITLE,
                 space_before=TITLE_BEFORE, space_after=TITLE_AFTER,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16)
    ensure_style(doc, "SectionTitle_Skills", size_pt=SZ_TITLE, bold=True, color=COLOR_TITLE,
                 space_before=TITLE_BEFORE, space_after=TITLE_AFTER,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16)
    ensure_style(doc, "SectionTitle_Awards", size_pt=SZ_TITLE, bold=True, color=COLOR_TITLE,
                 space_before=TITLE_BEFORE, space_after=TITLE_AFTER,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16)
    ensure_style(doc, "SectionTitle_Summary", size_pt=SZ_TITLE, bold=True, color=COLOR_TITLE,
                 space_before=TITLE_BEFORE, space_after=TITLE_AFTER,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT, line_pt=16)

    # ── 经历标题行（10.6pt 粗）──
    ensure_style(doc, "Education_ItemTitle", size_pt=SZ_BODY, bold=True, color=COLOR_BODY,
                 space_before=3, space_after=1, line_pt=LINE_BODY)
    ensure_style(doc, "Work_ItemTitle",      size_pt=SZ_BODY, bold=True, color=COLOR_BODY,
                 space_before=3, space_after=1, line_pt=LINE_BODY)
    ensure_style(doc, "Project_ItemTitle",   size_pt=SZ_BODY, bold=True, color=COLOR_BODY,
                 space_before=3, space_after=1, line_pt=LINE_BODY)

    # ── bullet / 正文（10.6pt 常规）──
    ensure_style(doc, "Education_Body", size_pt=SZ_BODY, bold=False, color=COLOR_BODY,
                 space_before=0, space_after=1, line_pt=LINE_BODY)
    ensure_style(doc, "Work_Bullet",    size_pt=SZ_BODY, bold=False, color=COLOR_BODY,
                 space_before=0, space_after=1, line_pt=LINE_BODY)
    ensure_style(doc, "Project_Bullet", size_pt=SZ_BODY, bold=False, color=COLOR_BODY,
                 space_before=0, space_after=1, line_pt=LINE_BODY)
    ensure_style(doc, "Summary_Bullet",  size_pt=SZ_BODY, bold=False, color=COLOR_BODY,
                 space_before=0, space_after=1, line_pt=LINE_BODY)

    # ── Skills / Awards ──
    ensure_style(doc, "Skill_Line", size_pt=SZ_BODY, bold=False, color=COLOR_BODY,
                 space_before=1, space_after=1, line_pt=LINE_BODY)
    ensure_style(doc, "Award_Line", size_pt=SZ_BODY, bold=False, color=COLOR_BODY,
                 space_before=1, space_after=1, line_pt=LINE_BODY)

    # ══════════════════════════════════════════════════════════════
    # 1. Profile 区（姓名 → 求职意向 → 联系方式，左对齐 + 右上角照片框）
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph(style="Profile_Name")
    _ph(p, "{{profile.name}}", size_pt=SZ_NAME, bold=True, color=COLOR_TITLE)

    # 求职意向（在联系方式之前，与 PDF 一致）
    p = doc.add_paragraph(style="Profile_Target")
    _run(p, "求职意向：", size_pt=SZ_CONTACT, bold=False, color=COLOR_BODY)
    _ph(p, "{{profile.target_position}}", size_pt=SZ_CONTACT, color=COLOR_BODY)

    # 联系方式：电话 丨 邮箱 丨 所在地
    p = doc.add_paragraph(style="Profile_Line")
    _run(p, "电话：", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _ph(p, "{{profile.phone}}", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _run(p, " 丨 ", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _run(p, "邮箱：", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _ph(p, "{{profile.email}}", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _run(p, " 丨 ", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _run(p, "所在地：", size_pt=SZ_CONTACT, color=COLOR_BODY)
    _ph(p, "{{profile.location}}", size_pt=SZ_CONTACT, color=COLOR_BODY)

    # 照片占位框（浮动右上角）
    _add_photo_placeholder(doc)

    # ══════════════════════════════════════════════════════════════
    # 2. 教育背景
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph("教育背景", style="SectionTitle_Education")
    _add_bottom_border(p)
    # ItemTitle: 时间 \t 学校 \t 专业（学历）
    _make_item_title(doc, "Education_ItemTitle", [
        ("edu.start_time", "-"),
        ("edu.end_time", "\t"),
        ("edu.school", "\t"),
        ("edu.major", "（"),
        ("edu.degree", "）"),
    ])
    # 主修课程 bullet
    _make_bullet(doc, "Education_Body", "edu.description")
    # GPA bullet（可选，空则自动删行）
    _make_bullet(doc, "Education_Body", "edu.gpa")

    # ══════════════════════════════════════════════════════════════
    # 3. 实习经历
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph("实习经历", style="SectionTitle_Work")
    _add_bottom_border(p)
    _make_item_title(doc, "Work_ItemTitle", [
        ("work.start_time", "-"),
        ("work.end_time", "\t"),
        ("work.company", "\t"),
        ("work.role", ""),
    ])
    _make_bullet(doc, "Work_Bullet", "work.bullet")

    # ══════════════════════════════════════════════════════════════
    # 4. 项目经历
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph("项目经历", style="SectionTitle_Project")
    _add_bottom_border(p)
    _make_item_title(doc, "Project_ItemTitle", [
        ("project.start_time", "-"),
        ("project.end_time", "\t"),
        ("project.name", "\t"),
        ("project.role", ""),
    ])
    _make_bullet(doc, "Project_Bullet", "project.bullet")

    # ══════════════════════════════════════════════════════════════
    # 5. 技能专长
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph("技能专长", style="SectionTitle_Skills")
    _add_bottom_border(p)
    # Skill_Line: 分类： 顿号分隔的技能项（分类加粗，技能项常规）
    p = doc.add_paragraph(style="Skill_Line")
    _ph(p, "{{skill.category}}", size_pt=SZ_BODY, bold=True, color=COLOR_BODY)
    _run(p, "：", size_pt=SZ_BODY, bold=True, color=COLOR_BODY)
    _ph(p, "{{skill.items}}", size_pt=SZ_BODY, bold=False, color=COLOR_BODY)

    # ══════════════════════════════════════════════════════════════
    # 6. 荣誉奖项（可选）
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph("荣誉奖项", style="SectionTitle_Awards")
    _add_bottom_border(p)
    p = doc.add_paragraph(style="Award_Line")
    _run(p, "⚫", size_pt=SZ_BODY, color=COLOR_BODY)
    _ph(p, "{{award}}", size_pt=SZ_BODY, color=COLOR_BODY)

    # ══════════════════════════════════════════════════════════════
    # 7. 自我评价（可选）
    # ══════════════════════════════════════════════════════════════
    p = doc.add_paragraph("自我评价", style="SectionTitle_Summary")
    _add_bottom_border(p)
    _make_bullet(doc, "Summary_Bullet", "summary.bullet")

    # ── 保存 ──
    doc.save(OUT_PATH)
    print(f"✓ 生成模板: {OUT_PATH}")
    verify(OUT_PATH)


def verify(path):
    from docx import Document as Doc
    doc = Doc(path)

    expected_styles = {
        "Profile_Name", "Profile_Target", "Profile_Line",
        "SectionTitle_Education", "SectionTitle_Work", "SectionTitle_Project",
        "SectionTitle_Skills", "SectionTitle_Awards", "SectionTitle_Summary",
        "Education_ItemTitle", "Education_Body",
        "Work_ItemTitle", "Work_Bullet",
        "Project_ItemTitle", "Project_Bullet",
        "Skill_Line", "Award_Line", "Summary_Bullet",
    }
    actual = {s.name for s in doc.styles}
    missing = expected_styles - actual
    if missing:
        raise SystemExit(f"模板缺少命名样式: {missing}")

    # 校验占位符独占 Run（R5）
    import re
    problems = []
    for p in doc.paragraphs:
        for r in p.runs:
            t = r.text or ""
            if "{{" in t or "}}" in t:
                stripped = t.strip()
                if not re.fullmatch(r'\{\{[a-zA-Z0-9_.]+\}\}', stripped):
                    problems.append(f"style={p.style.name} Run={t!r}（占位符未独占 Run）")

    # 校验无"含内容"的文本框（照片占位框 txbxContent 为空，允许）
    for tx in doc.element.body.findall('.//' + qn('w:txbxContent')):
        # 含文字段落的才报错
        for t_el in tx.findall('.//' + qn('w:t')):
            if (t_el.text or "").strip():
                problems.append("存在含文字的文本框（违反 R1）")
                break

    if problems:
        print("✗ 模板校验失败:")
        for m in problems:
            print("  -", m)
        raise SystemExit(1)
    print("✓ 模板校验通过（style 齐全 / 占位符独占 Run / 无内容文本框）")


if __name__ == "__main__":
    build()
