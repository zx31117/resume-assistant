"""Diagnose blank space in generated DOCX — check page breaks, section breaks, and spacing.

V1.4：默认从 runtime DOCX_OUTPUT_DIR 读取；兼容旧用法（第一个 CLI 参数传绝对路径时直接用）。
"""
import os
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))
from core.config import settings  # noqa: E402

from docx import Document  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402

if len(sys.argv) > 1:
    path = sys.argv[1]
else:
    path = os.path.join(settings.DOCX_OUTPUT_DIR, "resume_53ed02a6-3a45-4c53-b8a7-0a39b1bee71e_pm_template.docx")
print(f"[diag_docx] reading: {path}")
doc = Document(path)

print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total sections: {len(doc.sections)}")
print()

for i, p in enumerate(doc.paragraphs):
    style_name = p.style.name if p.style else "None"
    text = p.text[:100] if p.text else ""
    has_pb = False
    for run in p.runs:
        xml_str = run._r.xml
        if "br type=" in xml_str and "page" in xml_str:
            has_pb = True
            break
        if "lastRenderedPageBreak" in xml_str:
            has_pb = True
            break

    pf = p.paragraph_format
    space_before = pf.space_before.pt if pf.space_before else None
    space_after = pf.space_after.pt if pf.space_after else None

    pb_mark = " [PAGE_BREAK!]" if has_pb else ""
    text_repr = repr(text) if text else "(empty)"
    print(f"  [{i:2d}] style={style_name:<45s} sb={str(space_before):>6s}pt  sa={str(space_after):>6s}pt  text={text_repr}{pb_mark}")

print()
print("--- Section breaks ---")
for i, section in enumerate(doc.sections):
    print(f"  Section {i}: page_break_before={section.page_break_before}")
    print(f"    top_margin={section.top_margin}, bottom_margin={section.bottom_margin}")
