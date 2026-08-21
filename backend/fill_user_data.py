"""用虚构 mock 数据填充模板，验收最终效果。

本脚本仅供本地验收使用，不包含真实 PII。
所有数据均为虚构样例，用于验证渲染流程。

输出：${DOCX_OUTPUT_DIR}/resume_user_mock.docx
"""
import os
import sys
import json

BACKEND_ROOT = os.path.dirname(os.path.abspath(__file__))
if BACKEND_ROOT not in sys.path:
    sys.path.insert(0, BACKEND_ROOT)

# V1.4：输出统一走 runtime DOCX_OUTPUT_DIR，避免源码树被 .docx / .txt 等产物污染
from core.config import settings  # noqa: E402

OUTPUT_DIR = settings.DOCX_OUTPUT_DIR
os.makedirs(OUTPUT_DIR, exist_ok=True)


def build_user_resume():
    """用虚构 mock 数据构建 ResumeDocument。"""
    from models.resume_document import (
        Profile, EducationItem, WorkItem, ProjectItem,
        SkillGroup, ResumeDocument,
    )

    profile = Profile(
        name="张示例",
        phone="13800001111",
        email="zhangshili@example.com",
        location="示例市示例区",
        target_position="AI产品经理",
        summary="",  # mock 数据无自我评价章节，留空（可选章节不渲染）
    )

    education = [
        EducationItem(
            school="示例大学",
            major="计算机科学与技术",
            degree="本科",
            start_time="2023.09",
            end_time="2027.06",
            gpa="",  # mock 数据无 GPA 行
            description="主修课程：C语言程序设计、数据库、计算机组成原理、操作系统、软件工程等",
            priority=0.9,
        ),
    ]

    work = [
        WorkItem(
            company="示例科技有限公司",
            role="产品经理（实习）",
            start_time="2026.07",
            end_time="至今",
            bullets=[
                "需求管理：主动建立并维护需求池，梳理约40项产品需求，合并重复功能、调整模块归属，按照 P0（核心链路）—P3（体验优化）建立优先级，提升需求管理规范性及团队协作效率；",
                "流程优化：梳理 AI 批改作业、错题归纳、智能组卷等核心业务流程，重新规划页面跳转逻辑与功能入口，确保高频功能 2 次点击内完成访问，降低用户操作成本；",
                "产品设计：基于优化后的业务流程，使用墨刀输出高保真产品原型，同步完善页面交互、功能逻辑及异常流程设计，支撑需求评审与研发开发；",
                "PRD规范：结合敏捷开发模式重新整理 PRD，统一功能描述、流程图、优先级及验收标准，减少产品与研发间的信息偏差，提高需求沟通效率；",
                "协同推进：参与产品方案讨论，与产品经理、研发协同推进 AI 错题本功能迭代，产品预计于 2026 年 9 月上线，面向约 3000–5000 名用户。",
            ],
            priority=0.98,
        ),
        WorkItem(
            company="示例通讯有限公司",
            role="影像测评（实习）",
            start_time="2026.3",
            end_time="2026.6",
            bullets=[
                "影像性能测试：参与影像系统迭代版本的测试执行工作，按照标准化测试规范输出AE/AF性能、帧率表现及成像质量等客观数据，用于支持研发侧进行版本调测与优化评估；",
                "测试：在统一测试环境与流程约束下，完成多机型、多场景影像测试任务，确保测试数据采集过程符合要求，保证数据可复现与一致性；",
                "结果反馈与协同：统一测试环境将异常现象及测试数据反馈至调测与开发团队，支持后续问题定位。",
            ],
            priority=0.82,
        ),
    ]

    projects = [
        ProjectItem(
            name="语音交互状态反馈系统（氛围灯）",
            role="个人项目",
            start_time="2026.01",
            end_time="2026.03",
            bullets=[
                "需求动察：调研车载语音交互使用场景，发现连续监听状态下系统反馈依赖语音与屏幕，存在状态不可感知问题，用户难以判断是否仍处于可输入状态，导致自然对话被误触发为指令；",
                "产品设计：基于车载氛围灯设计状态反馈机制，将语音交互多阶段流程收敛为\"是否可输入\"二元模型，并设计对应状态映射（可输入：青色 / 不可输入：无灯或弱亮），补充语音与屏幕之外的持续状态通道；",
                "方案取舍：对比语音增强识别与意图判断方案，选择\"状态透明化\"路径，降低系统误判风险；同时收敛状态信息量，避免多状态映射带来的认知负担；",
                "原型验证：完成 Unity 交互 Demo 验证状态切换逻辑，并让 5 名用户小样本验证，初步验证可感知和弱干扰。",
            ],
            priority=0.9,
        ),
    ]

    skills = [
        SkillGroup(category="产品能力", items=["需求分析", "需求拆解", "产品流程设计", "PRD编写", "原型设计", "需求管理", "敏捷开发协作"]),
        SkillGroup(category="原型工具", items=["墨刀", "Axure", "Unity", "ProcessOn", "Office"]),
        SkillGroup(category="AI工具", items=["ChatGPT", "Claude", "TRAE"]),
    ]

    awards = []  # mock 数据无获奖章节

    doc = ResumeDocument(
        profile=profile,
        education=education,
        work=work,
        projects=projects,
        skills=skills,
        awards=awards,
        meta={
            "jd_position": "AI产品经理",
            "matched_count": 0,
            "total_experiences": 3,
            "user_id": "mock_user",
            "profile_source": "mock_data",
        },
    )
    doc = doc.to_standard()
    print(f"  ✓ profile: {profile.name} / {profile.target_position}")
    print(f"  ✓ education: {len(education)} 条, work: {len(work)} 条, projects: {len(projects)} 条")
    print(f"  ✓ skill_groups: {len(skills)} 组, awards: {len(awards)} 条")
    print()
    return doc


def main():
    from services.template_renderer import TemplateRenderer
    from services import layout_optimizer

    # Step 0: 构建模板
    print("=" * 68)
    print("Step 0: build pm_template.docx")
    print("=" * 68)
    import importlib.util
    build_script = os.path.join(BACKEND_ROOT, "templates", "_build_templates.py")
    spec = importlib.util.spec_from_file_location("_build_templates_mod", build_script)
    build_mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(build_mod)
    build_mod.build()
    print()

    # Step 1: 构建用户数据
    print("=" * 68)
    print("Step 1: build user ResumeDocument（张示例 - 虚构 mock 数据）")
    print("=" * 68)
    resume_doc = build_user_resume()

    # Step 2+3: 渲染 + 排版优化
    print("=" * 68)
    print("Step 2+3: TemplateRenderer.render → LayoutOptimizer.optimize")
    print("=" * 68)
    renderer = TemplateRenderer("pm_template", backend_root=BACKEND_ROOT)
    doc, render_warnings, _render_stats = renderer.render(resume_doc)
    print(f"  [Renderer] spec: {renderer.spec.id} v{renderer.spec.version}")
    print(f"  [Renderer] page_limit: {renderer.spec.layout.page_limit}")
    if render_warnings:
        print(f"  [Renderer] warnings ({len(render_warnings)}):")
        for w in render_warnings:
            print(f"    - {w}")
    else:
        print("  [Renderer] warnings: 0")

    page_limit = renderer.spec.layout.page_limit
    est_before = layout_optimizer.estimate_pages(doc)
    applied_rules, capacity_warnings = layout_optimizer.optimize(doc, page_limit=page_limit)
    est_after = layout_optimizer.estimate_pages(doc)
    print(f"  [Layout]   pages (estimate): {est_before} → {est_after}  （target ≤ {page_limit}）")
    if applied_rules:
        print(f"  [Layout]   optimizations applied ({len(applied_rules)}):")
        for r in applied_rules:
            print(f"    - {r}")
    else:
        print("  [Layout]   optimizations applied: 0")
    if capacity_warnings:
        print("  [Layout]   capacity warnings:")
        for w in capacity_warnings:
            print(f"    - {w}")
    print()

    # Step 4: 保存
    print("=" * 68)
    print("Step 4: save DOCX")
    print("=" * 68)
    file_name = "resume_user_mock.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    doc.save(file_path)
    size_kb = os.path.getsize(file_path) / 1024.0
    print(f"  ✓ 保存成功: {file_path}")
    print(f"    文件大小: {size_kb:.1f} KB")
    print()

    # Step 5: 内容验收报告（提取生成内容逐项核对）
    print("=" * 68)
    print("Step 5: 内容验收报告（生成内容逐项核对）")
    print("=" * 68)
    from docx import Document as Doc
    chk = Doc(file_path)
    all_text = "\n".join(p.text for p in chk.paragraphs if p.text.strip())

    checks = []
    # 关键内容核对（mock 数据）
    key_items = [
        ("姓名", "张示例"),
        ("电话", "13800001111"),
        ("邮箱", "zhangshili@example.com"),
        ("所在地", "示例市示例区"),
        ("求职意向", "AI产品经理"),
        ("学校", "示例大学"),
        ("专业", "计算机科学与技术"),
        ("公司1", "示例科技有限公司"),
        ("公司2", "示例通讯有限公司"),
        ("项目名", "语音交互状态反馈系统"),
        ("技能1", "产品能力"),
        ("技能2", "原型工具"),
        ("技能3", "AI工具"),
    ]
    for label, kw in key_items:
        ok = kw in all_text
        checks.append((f"内容含「{label}」（{kw}）", ok))

    # 章节标题核对
    for title in ["教育背景", "实习经历", "项目经历", "技能专长"]:
        ok = title in all_text
        checks.append((f"章节标题「{title}」", ok))

    # 可选章节不应出现（mock 数据无）
    for absent in ["荣誉奖项", "自我评价"]:
        ok = absent not in all_text
        checks.append((f"无「{absent}」章节（mock 数据无此项）", ok))

    # bullet ⚫ 核对
    bullet_count = sum(1 for p in chk.paragraphs if (p.runs and "⚫" in (p.runs[0].text or "")))
    checks.append((f"⚫ bullet 数量 = {bullet_count}（预期 1+5+3+4=13）", bullet_count == 13))

    # 照片框
    from docx.oxml.ns import qn
    has_photo = False
    for p in chk.paragraphs:
        for r in p.runs:
            if r._element.findall(qn('w:drawing')):
                has_photo = True
    checks.append(("右上角照片占位框存在", has_photo))

    all_pass = True
    for desc, ok in checks:
        tag = "PASS" if ok else "FAIL"
        print(f"  [{tag}] {desc}")
        all_pass = all_pass and ok

    print()
    print(f"页数估算: {est_after}（目标 ≤ {page_limit}）")
    print(f"文件: {file_path}")
    if all_pass:
        print("结论: ✓ 验收通过，内容与 mock 数据一致")
    else:
        print("结论: ✗ 存在差异项，请检查")

    # 保存验收报告
    report_path = os.path.join(OUTPUT_DIR, "验收报告_用户数据填充.txt")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("=== 用户数据填充验收报告（张示例 - mock 数据）===\n")
        f.write(f"数据来源: 虚构 mock 数据（无真实 PII）\n")
        f.write(f"输出文件: {file_path}\n")
        f.write(f"页数估算: {est_after}（目标 ≤ {page_limit}）\n\n")
        for desc, ok in checks:
            f.write(f"[{'PASS' if ok else 'FAIL'}] {desc}\n")
        f.write(f"\n结论: {'✓ 验收通过' if all_pass else '✗ 存在差异'}\n")
    print(f"  报告已保存: {report_path}")


if __name__ == "__main__":
    main()
