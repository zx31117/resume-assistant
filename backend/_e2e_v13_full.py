"""V1.3 完整流程 E2E：
input/简历.pdf（用户上传简历）→ 结构化 Profile + Experiences
input/JD.txt（岗位 JD，GBK 编码）
→ 写入 SQL（建 User + 建 Experience + VectorIndexJob UPSERT → 同步 Chroma）
→ V1.3 核心 ResumeGenerationService.generate_docx
→ 输出 DOCX + 验收信息

运行目录：backend/
需环境变量：至少有可工作的 OPENAI_API_KEY（或 ARK_API_KEY，视 .env 配置）
"""
from __future__ import annotations

import json
import os
import sys
from pathlib import Path

BACKEND_ROOT = Path(__file__).resolve().parent
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))
WORKTREE_ROOT = BACKEND_ROOT.parent
INPUT_DIR = WORKTREE_ROOT / "input"

# V1.4：DOCX 生成结果统一走 runtime output（源码树保持干净，避免 output/ 被误提交）
from core.config import settings  # noqa: E402

OUTPUT_DIR = Path(settings.DOCX_OUTPUT_DIR)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------- #
# 1) 读文件：简历.pdf + JD.txt
# ---------------------------------------------------------------------- #
def step1_read_inputs() -> tuple[str, bytes]:
    print("=" * 68)
    print("Step 1: 读取 input/简历.pdf 和 input/JD.txt（JD 为 GBK 编码）")
    print("=" * 68)

    jd_path = INPUT_DIR / "JD.txt"
    pdf_path = INPUT_DIR / "简历.pdf"
    assert jd_path.exists(), f"找不到 {jd_path}"
    assert pdf_path.exists(), f"找不到 {pdf_path}"

    # JD 用 GBK（ANSI）读，若失败再试 UTF-8 fallback
    try:
        jd_text = jd_path.read_text(encoding="gbk")
        print("  ✓ JD.txt 以 GBK 读取成功")
    except UnicodeDecodeError:
        jd_text = jd_path.read_text(encoding="utf-8")
        print("  ✓ JD.txt 以 UTF-8 读取成功（GBK 失败 fallback）")

    pdf_bytes = pdf_path.read_bytes()
    print(f"  ✓ 简历.pdf 读取 {len(pdf_bytes):,} bytes")
    print(f"    JD 长度: {len(jd_text)} 字符")
    print()
    return jd_text, pdf_bytes


# ---------------------------------------------------------------------- #
# 2) PDF → 文本；然后抽取 Profile + Experiences
# ---------------------------------------------------------------------- #
def step2_parse_and_extract(pdf_bytes: bytes):
    from services import resume_parser, experience_extractor, llm_service
    from api.schemas import RequestProfile

    print("=" * 68)
    print("Step 2: PDF → 纯文本 → LLM 抽取 Profile + Experiences")
    print("=" * 68)

    pdf_text = resume_parser.parse_pdf(pdf_bytes)
    print(f"  ✓ PDF 文本提取完成，长度 {len(pdf_text)} 字符")
    if len(pdf_text) < 100:
        print(f"  ⚠ PDF 文本过短，可能是扫描件或纯图片。先继续看 LLM 能否保底抽。")
    print()

    # 2a) 抽 Profile（姓名/电话/邮箱/所在地/目标岗位/summary）——强类型 strict
    profile_prompt_sys = (
        "你是资深简历解析专家。只从简历原文中抽取用户基本信息。"
        "严格基于原文，禁止编造未提及的字段。找不到就填空字符串。"
    )
    profile_prompt_user = (
        "从以下简历全文抽取基本信息，只输出 JSON，不要解释，不得加 Markdown 围栏。\n"
        "字段：name（姓名）, phone（手机号/电话）, email（邮箱）, location（所在地/城市）, "
        "target_position（求职意向/目标岗位）, summary（自我评价摘要，2-4句；没有就填空字符串）。\n\n"
        f"简历全文：\n{pdf_text[:6000]}\n"
    )
    profile: RequestProfile = llm_service.chat_structured(
        profile_prompt_sys, profile_prompt_user,
        schema=RequestProfile,
        default=RequestProfile(name=""),
        strict=False,  # 解析结果允许部分字段空（由 Builder 兜底），但强类型结构必须对
    )
    print("  ✓ Profile 抽取:")
    print(f"    - 姓名: {profile.name!r}")
    print(f"    - 电话: {profile.phone!r}")
    print(f"    - 邮箱: {profile.email!r}")
    print(f"    - 所在地: {profile.location!r}")
    print(f"    - 目标岗位: {profile.target_position!r}")
    print(f"    - summary 长度: {len(profile.summary or '')} 字符")
    print()

    # 2b) 抽 Experiences（复用 V1.1 现有模块）
    exps = experience_extractor.extract_experiences(pdf_text)
    print(f"  ✓ Experiences 抽取 {len(exps)} 条（去重后）")
    for i, e in enumerate(exps, 1):
        title = e.get("title") or e.get("role") or ""
        company = e.get("company") or ""
        ttype = e.get("type") or ""
        time = e.get("time") or ""
        print(f"    {i}. [{ttype}] {title} @ {company}  ({time})  "
              f"skills={len(e.get('skills') or [])}  bullets/desc_len={len(e.get('description') or e.get('raw_text') or '')}")
    print()
    return pdf_text, profile, exps


# ---------------------------------------------------------------------- #
# 3) 写入 SQL：建表 + 建 User + 每条经历 create_experience（触发 Job）
# ---------------------------------------------------------------------- #
def step3_persist(profile, extracted_exps: list[dict]) -> tuple[object, str]:
    from database.init_db import init_db
    from database.session import SessionLocal
    from database import models
    from services import experience_service
    from core.config import settings

    print("=" * 68)
    print("Step 3: init_db → 建 User → 逐条 create_experience（同事务 Job + 同步向量）")
    print("=" * 68)

    init_db()
    db = SessionLocal()

    # 3a) 建/取 User（优先按 email 去重）
    user = None
    email = (profile.email or "").strip()
    if email:
        user = db.query(models.User).filter(models.User.email == email).first()
    if user is None:
        user = models.User(
            name=(profile.name or "").strip() or None,
            email=email or None,
        )
        db.add(user)
        db.commit()
        db.refresh(user)
    user_id = user.id
    print(f"  ✓ User: id={user_id} name={user.name!r} email={user.email!r}")

    # 3b) 逐条写入 Experience，触发 VectorIndexJob UPSERT + 同步执行
    ok_cnt = 0
    failed_cnt = 0
    for i, e in enumerate(extracted_exps, 1):
        data = {
            "type": e.get("type", "") or "",
            "title": e.get("title", "") or "",
            "company": e.get("company", "") or "",
            "time": e.get("time", "") or "",
            "role": e.get("role", "") or "",
            "description": e.get("description", "") or "",
            "skills": list(e.get("skills") or []) or [],
            "achievements": list(e.get("achievements") or []) or [],
            "raw_text": e.get("raw_text", "") or "",
        }
        try:
            exp = experience_service.create_experience(db, user_id, data)
            print(f"  ✓ 经历 {i}/{len(extracted_exps)} 落库 + 索引成功: exp_id={exp.id} vector_id={exp.vector_id!r}")
            ok_cnt += 1
        except Exception as ex:
            print(f"  ✗ 经历 {i}/{len(extracted_exps)} 失败: {type(ex).__name__}: {ex!r}")
            failed_cnt += 1

    # 3c) 检查是否有 FAILED / PENDING 的 Job
    from services import vector_index_sync
    try:
        stats = vector_index_sync.ensure_user_index_ready(db, user_id)
        print(f"  ✓ ensure_user_index_ready -> {json.dumps(stats, ensure_ascii=False)}")
    except Exception as e:
        # VectorIndexNotReadyError 会含 failed_ids / pending_ids；先打印，仍然继续（generate-docx 阶段会再次检查）
        print(f"  ⚠ ensure_user_index_ready 异常（下面 generate-docx 会再检查）：{type(e).__name__}: {e!r}")

    print(f"  汇总: {ok_cnt} 条成功 / {failed_cnt} 条失败")
    print()
    return db, user_id


# ---------------------------------------------------------------------- #
# 4) V1.3 核心：generate-docx
# ---------------------------------------------------------------------- #
def step4_generate_docx(db, user_id: str, jd_text: str, profile):
    from api.schemas import ResumeDocxGenerateRequest
    from services import resume_generation_service

    print("=" * 68)
    print("Step 4: V1.3 核心 — generate_docx (8 stages)")
    print("=" * 68)

    req = ResumeDocxGenerateRequest(
        user_id=user_id,
        template_id="pm_template",
        jd_text=jd_text,
        profile=profile,
        top_k=5,
    )
    print(f"  请求: user_id={req.user_id} template_id={req.template_id} top_k={req.top_k}")
    print(f"         profile.name={profile.name!r} target_position={profile.target_position!r}")
    print(f"         JD 长度={len(jd_text)} 字符")
    print()

    resp = resume_generation_service.generate_docx(db, req)
    print(f"  ✓ 成功!")
    print(f"    - file_name: {resp.file_name}")
    print(f"    - file_path: {resp.file_path}")
    print(f"    - download_url: {resp.download_url}")
    print(f"    - page_count: {resp.page_count}")
    print(f"    - profile_source: {resp.profile_source}")
    print(f"    - build_counts: {json.dumps(resp.build_counts or {}, ensure_ascii=False)}")
    print()
    print(f"  Stages:")
    for s in resp.stages:
        status_symbol = {"done": "✓", "failed": "✗", "running": "→"}.get(s.status, "?")
        print(f"    {status_symbol} {s.stage:<22}  {s.status:<7} {s.duration_ms:>7} ms   {s.note or ''}")
    print()
    print(f"  Matched ids ({len(resp.matched_experience_ids)}): {resp.matched_experience_ids}")
    print(f"  Rendered ids ({len(resp.rendered_experience_ids)}): {resp.rendered_experience_ids}")
    if resp.warnings:
        print(f"  Warnings ({len(resp.warnings)}):")
        for w in resp.warnings:
            print(f"    - {w}")
    else:
        print("  Warnings: 0")
    print()
    return resp


# ---------------------------------------------------------------------- #
# 5) 验收：读取生成的 DOCX，核对关键内容
# ---------------------------------------------------------------------- #
def step5_acceptance(resp, profile):
    from docx import Document as Doc
    from docx.oxml.ns import qn

    print("=" * 68)
    print("Step 5: 生成内容验收（基于最终 DOCX）")
    print("=" * 68)

    file_path = BACKEND_ROOT / resp.file_path
    assert file_path.exists(), f"DOCX 不存在: {file_path}"

    doc = Doc(str(file_path))
    all_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    # 关键核对：
    checks = []
    def _add(label, cond, detail=""):
        checks.append((label, bool(cond), detail))

    _add(f"姓名 「{profile.name}」 出现", profile.name and profile.name in all_text, profile.name or "")
    if profile.phone:
        _add(f"电话 「{profile.phone}」 出现", profile.phone in all_text, profile.phone)
    if profile.email:
        _add(f"邮箱 「{profile.email}」 出现", profile.email in all_text, profile.email)
    if profile.location:
        _add(f"所在地 「{profile.location}」 出现（若解析为空则本项自动跳过）", True if not profile.location else profile.location in all_text)
    # 章节标题（教育背景/实习经历/工作经历/项目经历/技能专长 至少要出现一些）
    titles_seen = []
    for t in ["教育背景", "实习经历", "工作经历", "项目经历", "技能专长"]:
        if t in all_text:
            titles_seen.append(t)
    _add(f"至少出现 2 个核心章节标题（实际：{titles_seen}）", len(titles_seen) >= 2, str(titles_seen))
    # Bullet
    bullet_count = sum(1 for p in doc.paragraphs if (p.runs and "⚫" in (p.runs[0].text or "")))
    _add(f"⚫ bullet 数 ≥ 1（实际 {bullet_count}）", bullet_count >= 1, str(bullet_count))
    # 照片占位框
    has_photo = False
    for p in doc.paragraphs:
        for r in p.runs:
            if r._element.findall(qn("w:drawing")):
                has_photo = True
                break
    _add("右上角照片占位框存在", has_photo)

    # 占位符扫描
    import re
    unreplaced: list[str] = []
    for pat in [r"\{\{([^{}]+)\}\}", r"\[\[([^\[\]]+)\]\]"]:
        for m in re.finditer(pat, all_text):
            unreplaced.append(m.group(0))
    uniq_unreplaced = sorted(set(unreplaced))
    _add(f"无未替换占位符（实际命中 {len(uniq_unreplaced)} 样例={uniq_unreplaced[:5]}）", len(uniq_unreplaced) == 0)

    all_pass = True
    for desc, ok, detail in checks:
        print(f"  [{'PASS' if ok else 'FAIL'}] {desc}{(' — ' + detail) if detail else ''}")
        all_pass = all_pass and ok

    print()
    print(f"  DOCX 路径: {file_path}")
    print(f"  大小: {file_path.stat().st_size / 1024:.1f} KB")
    print(f"  页数(estimate): {resp.page_count}")
    print(f"  最终结论: {'✓ 内容级验收通过' if all_pass else '✗ 存在差异项'}")
    print()
    report_path = OUTPUT_DIR / "验收报告_V1.3_完整流程.txt"
    with report_path.open("w", encoding="utf-8") as f:
        f.write("=== V1.3 完整流程验收报告 ===\n")
        f.write(f"输入: input/简历.pdf + input/JD.txt\n")
        f.write(f"输出 DOCX: {file_path}\n")
        f.write(f"页数(estimate): {resp.page_count}\n")
        f.write(f"profile_source: {resp.profile_source}\n")
        f.write(f"matched_experience_ids: {resp.matched_experience_ids}\n")
        f.write(f"rendered_experience_ids: {resp.rendered_experience_ids}\n\n")
        f.write("--- Stages ---\n")
        for s in resp.stages:
            f.write(f"{s.stage:<22}  {s.status:<7} {s.duration_ms:>7} ms   {s.note or ''}\n")
        f.write("\n--- Warnings ---\n")
        for w in resp.warnings:
            f.write(f"- {w}\n")
        f.write("\n--- 内容核对 ---\n")
        for desc, ok, detail in checks:
            f.write(f"[{'PASS' if ok else 'FAIL'}] {desc}{(' — ' + detail) if detail else ''}\n")
    print(f"  ✓ 报告已保存: {report_path}")


def main():
    # print .env key 状态（避免泄露，只打印是否配置）
    from core.config import settings
    ark_set = bool(settings.ARK_API_KEY)
    openai_set = bool(os.getenv("OPENAI_API_KEY"))
    print(f"[env] ARK_API_KEY set? {ark_set}  /  OPENAI_API_KEY set? {openai_set}")
    print(f"[env] LLM_MODEL={settings.LLM_MODEL!r}  EMBEDDING_MODEL={settings.EMBEDDING_MODEL!r}")
    if not ark_set and not openai_set:
        print("⚠ 没有配置 ARK_API_KEY 也没配置 OPENAI_API_KEY，LLM 会立刻失败。"
              "请先在 backend/.env 或 backend/.env.local 中配置至少一个。")
        sys.exit(2)
    print()

    jd_text, pdf_bytes = step1_read_inputs()
    pdf_text, profile, exps = step2_parse_and_extract(pdf_bytes)
    db, user_id = step3_persist(profile, exps)
    try:
        resp = step4_generate_docx(db, user_id, jd_text, profile)
    finally:
        db.close()
    step5_acceptance(resp, profile)
    print("== V1.3 E2E 完成 ==")


if __name__ == "__main__":
    main()
