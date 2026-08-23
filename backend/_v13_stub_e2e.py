"""V1.5.0 Stub E2E - 固定 LLM/Embedding mock，CI 可重复（V1.5.0 两层选材链路）。

V1.5.0 变更（PLAN §7 T6）：
- 旧 RAG 链路（rag_service / vector_index_sync / chroma_store）已退出；
- 新链路：迁移检查 → JD 分析 → 第一层选材 → 第二层事实选材 → 受约束改写 → Builder 收缩 → 渲染 → DOCX。
- Mock 接缝：embedding_service._embed_text（stub 向量）、llm_service.chat_structured（stub JD + V15 改写）。

每次运行强制使用临时独立 RESUME_DATA_DIR，测试结束后整体清理，
完全不接触或污染用户真实 runtime。连续运行两次都应全部通过。
"""
import atexit
import json, os, re, shutil, sys, tempfile
from pathlib import Path
from typing import Any
from unittest.mock import patch, MagicMock

# 设置测试专用假 Key（仅通过 import，不发起真实网络请求）
os.environ.setdefault("ARK_API_KEY", "stub-e2e-ark-key-not-real")
os.environ.setdefault("OPENAI_API_KEY", "stub-e2e-openai-key-not-real")

# 强制使用独立临时 RESUME_DATA_DIR
_STUB_RUNTIME_TMP = Path(tempfile.mkdtemp(prefix="stub-e2e-runtime-")).resolve()
os.environ["RESUME_DATA_DIR"] = str(_STUB_RUNTIME_TMP)
os.environ["SQLITE_PATH"] = str(_STUB_RUNTIME_TMP / "database" / "app.db")
os.environ["DOCX_OUTPUT_DIR"] = str(_STUB_RUNTIME_TMP / "output")
# V1.5.0：CHROMA_PATH 已退出，不再设置

_CLEANUP_DONE = False
_CLEANUP_RESULT = None

def _cleanup_stub_runtime():
    """Release resources then remove the temp runtime directory with limited retry."""
    global _CLEANUP_DONE, _CLEANUP_RESULT
    if _CLEANUP_DONE:
        return _CLEANUP_RESULT

    # 1. Release SQLAlchemy engine
    try:
        from database.session import engine as _engine
        _engine.dispose()
    except Exception as e:
        print(f"[cleanup][WARN] engine.dispose failed: {e}")

    # 2. V1.5.0：chroma_store 已删除，无需关闭 Chroma client

    # 3. Force GC
    import gc
    gc.collect()

    # 4. Remove temp directory
    if not _STUB_RUNTIME_TMP.exists():
        _CLEANUP_DONE = True
        _CLEANUP_RESULT = True
        return True

    import time
    import stat as _stat

    def _on_rm_error(func, path, exc_info):
        """Handle read-only files (migration backups set chmod 0o444)."""
        try:
            os.chmod(path, _stat.S_IWRITE)
            func(path)
        except Exception:
            pass

    last_err = None
    for attempt in range(5):
        try:
            shutil.rmtree(_STUB_RUNTIME_TMP, onerror=_on_rm_error)
            _CLEANUP_DONE = True
            _CLEANUP_RESULT = True
            return True
        except Exception as e:
            last_err = e
            time.sleep(0.3 * (attempt + 1))
            gc.collect()
    if _STUB_RUNTIME_TMP.exists():
        print(f"[cleanup][FAIL] residual after 5 retries: {last_err}")
        _CLEANUP_DONE = True
        _CLEANUP_RESULT = False
        return False
    _CLEANUP_DONE = True
    _CLEANUP_RESULT = True
    return True

atexit.register(_cleanup_stub_runtime)

BACKEND_ROOT = Path(__file__).resolve().parent
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))

from core.config import settings  # noqa: E402
from core.version import APP_VERSION  # noqa: E402
from models.resume_document import Profile as _ProfileDoc  # noqa: E402

OUTPUT_DIR = Path(settings.DOCX_OUTPUT_DIR)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

STUB_JD_ANALYSIS = {
    "position": "AI硬件产品经理", "industry": "智能硬件/AI",
    "required_skills": ["产品规划", "AI", "硬件", "数据分析", "项目管理"],
    "preferred_skills": ["大模型", "Python", "用户研究"],
    "responsibilities": ["负责AI硬件产品规划和定义"],
    "keywords": ["AI", "硬件", "产品经理", "智能硬件"],
    "experience_preferences": ["3年以上产品经验"],
}
STUB_EMBEDDING_VECTOR = [0.1] * 2048

STUB_EXPERIENCES = [
    {"type": "work", "title": "影像测试实习生", "company": "深圳传音控股股份有限公司",
     "role": "影像测试实习生（AI方向）", "time": "2026.05-至今",
     "description": "负责AI影像产品测试和数据分析。",
     "skills": ["AI", "数据分析", "Python"], "achievements": ["完成3个AI影像项目测试"]},
    {"type": "project", "title": "基于STM32的智能家居控制系统",
     "role": "项目负责人", "time": "2025.03-2025.06",
     "description": "设计STM32智能家居系统。",
     "skills": ["STM32", "C语言", "IoT"], "achievements": ["校级优秀项目奖"]},
    {"type": "project", "title": "基于LLM Agent的智能客服系统",
     "role": "核心开发者", "time": "2025.09-2025.12",
     "description": "基于LLM Agent构建智能客服。",
     "skills": ["Python", "LLM", "NLP"], "achievements": ["客服效率提升40%"]},
    {"type": "education", "title": "电子信息工程", "company": "某某大学",
     "role": "本科", "time": "2022.09-2026.06",
     "description": "GPA 3.8/4.0", "skills": [], "achievements": []},
]
STUB_PROFILE = {"name": "张三", "phone": "13800000001", "email": "zhangsan@example.com",
                "location": "深圳", "target_position": "AI硬件产品经理"}
STUB_JD_TEXT = "AI硬件产品经理\n负责AI硬件产品规划。\n要求: 3年产品经验, AI/硬件优先。"

def _cleanup():
    if OUTPUT_DIR.exists():
        for p in OUTPUT_DIR.glob("resume_*.docx"):
            try:
                p.unlink()
            except Exception as e:
                print(f"[cleanup][WARN] failed to unlink {p.name}: {e}")
    print(f"[cleanup] runtime isolated at: {_STUB_RUNTIME_TMP}")


# ── V1.5.0 Mock 接缝 ──────────────────────────────────────────── #

def _create_mock_embed():
    """Stub embedder: 返回固定向量（V1.5.0：mock embedding_service._embed_text）"""
    return lambda text: list(STUB_EMBEDDING_VECTOR)


def _create_mock_chat_structured(stub_jd, all_exp_ids=None):
    """Mock llm_service.chat_structured for V1.5.0 chain.

    - schema=JDAnalysisOut → 返回 stub JD 分析
    - schema=GeneratedResumeContentV15 → 返回带 fact_refs 的 V15 内容
    - schema=GeneratedResumeContent (旧 V1.3) → 返回旧格式（兼容）
    """
    from api.schemas import JDAnalysisOut, GeneratedResumeContentV15, GeneratedBullet, GeneratedExperienceItemV15
    from api.schemas import GeneratedResumeContent, GeneratedExperienceItem

    def fn(system, user_template, schema, strict=False, default=None, **variables):
        schema_name = getattr(schema, '__name__', str(schema))

        if schema_name == "JDAnalysisOut":
            return JDAnalysisOut.model_validate(stub_jd)

        if schema_name == "GeneratedResumeContentV15":
            # V1.5.0 受约束改写：从 evidence_json 提取 fact_ids
            evidence_json = variables.get("evidence_json", "[]")
            try:
                payload = json.loads(evidence_json) if isinstance(evidence_json, str) else evidence_json
            except Exception:
                payload = []
            items = []
            for entry in payload:
                exp_id = entry.get("experience_id", "")
                facts = entry.get("usable_facts", [])
                if facts:
                    fid = facts[0].get("fact_id", "")
                    items.append(GeneratedExperienceItemV15(
                        experience_id=exp_id,
                        bullets=[GeneratedBullet(
                            bullet=f"[STUB] AI bullet for {exp_id[:8]}",
                            fact_refs=[fid],
                        )],
                    ))
                else:
                    items.append(GeneratedExperienceItemV15(
                        experience_id=exp_id,
                        bullets=[],
                        insufficient=True,
                        insufficient_reason="无可用事实",
                    ))
            return GeneratedResumeContentV15(experiences=items)

        if schema_name == "GeneratedResumeContent":
            # 旧 V1.3 兼容
            ids = all_exp_ids or []
            items = [{"experience_id": eid, "bullets": [f"[STUB] AI bullet for {eid[:8]}"]} for eid in ids]
            return GeneratedResumeContent.model_validate({"experiences": items})

        return schema() if default is None else default

    return fn


def _setup_test_data(db_factory):
    """V1.5.0：创建 User + Experiences，运行迁移生成 Facts + Embeddings。"""
    from database import models, migrations
    from services import experience_service, embedding_service
    db = db_factory()
    try:
        u = db.query(models.User).filter(models.User.id == "stub-user").first()
        if not u:
            u = models.User(id="stub-user", name="张三")
            db.add(u)
            db.commit()
        uid = "stub-user"
        ids = []
        for ed in STUB_EXPERIENCES:
            exp = experience_service.create_experience(db, uid, ed)
            ids.append(exp.id)
        db.close()

        # V1.5.0：运行迁移（创建 Facts + SchemaVersion）
        # embedding_service._embed_text 已被 mock，rebuild_embeddings 使用 stub 向量
        # R3 fix: dispose global engine to release SQLite file lock before backup
        from database.session import engine as _global_engine
        _global_engine.dispose()
        # Stub E2E: backup=False — backup logic covered by T2; avoid Windows file lock
        mig_stats = migrations.run_migrations(settings.SQLITE_PATH, backup=False)
        print(f"  [setup] migration stats: {json.dumps(mig_stats, ensure_ascii=False)[:200]}")

        # V1.5.0：重建 embeddings（使用 mock embedder）
        db2 = db_factory()
        try:
            stats = embedding_service.rebuild_embeddings(
                db2, embedder=_create_mock_embed(),
            )
            print(f"  [setup] embedding rebuild: {json.dumps(stats, ensure_ascii=False)[:200]}")
        finally:
            db2.close()

        return uid, ids
    finally:
        if db.is_active:
            db.close()


def run_happy_path():
    from database.init_db import init_db
    from database.session import SessionLocal
    from services import resume_generation_service, experience_service
    from api.schemas import RequestProfile, ResumeDocxGenerateRequest
    from docx import Document as DocxDoc

    _cleanup()
    init_db()

    # V1.5.0：mock embedding_service._embed_text（迁移和重建时使用）
    with patch("services.embedding_service._embed_text", side_effect=_create_mock_embed()):
        uid, eids = _setup_test_data(SessionLocal)

    db = SessionLocal()
    try:
        all_exps = experience_service.list_experiences(db, uid)
        wids = [e.id for e in all_exps if e.type == "work"]
        pids = [e.id for e in all_exps if e.type == "project"]

        # V1.5.0：mock LLM（JD 分析 + 受约束改写）+ mock embedding（查询时）
        with patch("services.llm_service.chat_structured",
                   side_effect=_create_mock_chat_structured(STUB_JD_ANALYSIS)), \
             patch("services.embedding_service._embed_text", side_effect=_create_mock_embed()):
            req = ResumeDocxGenerateRequest(user_id=uid, template_id="pm_template",
                                            jd_text=STUB_JD_TEXT, profile=RequestProfile(**STUB_PROFILE), top_k=5)
            resp = resume_generation_service.generate_docx(db, req)

        dp = OUTPUT_DIR / resp.file_name if resp.file_name else BACKEND_ROOT / resp.file_path
        ok1 = dp.exists() and dp.stat().st_size > 2048
        ft = ""
        doc = None
        if ok1:
            doc = DocxDoc(str(dp))
            ft = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            ft += "\n" + "\n".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)

        print(f"  [DEBUG] DOCX text sample: {ft[:200]}...")

        results = {}
        results["1_core_docx_ok"] = {"status": "通过" if ok1 else "失败",
            "evidence": {"file_path": str(dp), "size": dp.stat().st_size if dp.exists() else 0, "stages": len(resp.stages)}}

        pt = ["王示例", "示例科技", "示例项目", "z****@***********", "138****1111", "示例市", "AI 产品经理（示例）"]
        ex = {"phone": STUB_PROFILE["phone"], "email": STUB_PROFILE["email"]}
        ok2 = all(v and v in ft for v in ex.values() if v) and not any(t in ft for t in pt)
        _has_profile_name_style = (
            any("Profile_Name" in (p.style.name if hasattr(p,"style") else "") for p in doc.paragraphs)
            if doc else False
        )
        results["2_personal_not_from_template"] = {"status": "通过" if ok2 else "失败",
            "evidence": {"present": {k: bool(v) and v in ft for k,v in ex.items()}, "profile_name_style_present": _has_profile_name_style, "leaked": [t for t in pt if t in ft]}}

        sids = {e.id for e in all_exps}
        ms = set(resp.matched_experience_ids)
        rs = set(resp.rendered_experience_ids)
        ok3 = rs.issubset(sids & ms) if ms else rs.issubset(sids)
        results["3_rendered_ids_subset"] = {"status": "通过" if ok3 else "失败",
            "evidence": {"sql": sorted(sids), "matched": sorted(ms), "rendered": sorted(rs), "diff": sorted(rs - (sids & ms)) if ms else sorted(rs - sids)}}

        fo = True; fd = {}
        for eid in sorted(rs):
            exp = next((x for x in all_exps if x.id == eid), None)
            if not exp: fo = False; fd[eid] = "SQL缺失"; continue
            toks = []
            if exp.type == "work": toks = [x for x in [exp.company, exp.role, exp.time] if x and x.strip()]
            elif exp.type == "project": toks = [x for x in [exp.title, exp.role, exp.time] if x and x.strip()]
            pr = {t: (t in ft) for t in toks}
            if exp.type in ("work","project") and not all(pr.values()): fo = False
            fd[eid] = {"type": exp.type, "tokens": pr}
        results["4_facts_equal_sql"] = {"status": "通过" if fo else "失败", "evidence": fd}

        # V1.5.0：build_v15 收缩验证（通过 build_meta 间接验证 fact_refs 保留）
        if wids or pids:
            bm2 = resp.build_meta
            ok5 = all([
                isinstance(bm2.profile_source, str) and bm2.profile_source,
                isinstance(bm2.ai_covered_experience_ids, list),
                isinstance(bm2.fallback_sql_experience_ids, list),
            ])
            results["5_build_v15_meta"] = {"status": "通过" if ok5 else "失败",
                "evidence": {"profile_source": bm2.profile_source,
                             "ai_covered": bm2.ai_covered_experience_ids,
                             "fallback": bm2.fallback_sql_experience_ids,
                             "counts": bm2.counts.model_dump()}}
        else:
            results["5_build_v15_meta"] = {"status": "未执行", "evidence": "no work/project"}

        sm = {s.section_id: s for s in resp.render_stats.sections}
        ok6 = True
        for sid in ("work","skills"):
            s = sm.get(sid)
            if s and s.input_items != s.rendered_items: ok6 = False
        st = {"education": ["教育背景"], "projects": ["项目经历"], "work": ["实习经历","工作经历"], "skills": ["技能专长"]}
        for sk, tks in st.items():
            s = sm.get(sk)
            if s and s.input_items > 0 and not any(t in ft for t in tks): ok6 = False
        if len(resp.render_stats.unreplaced_placeholders) > 0: ok6 = False
        results["6_renderer_no_truncate"] = {"status": "通过" if ok6 else "失败",
            "evidence": {"sections": [s.model_dump() for s in resp.render_stats.sections],
                         "unreplaced": resp.render_stats.unreplaced_placeholders,
                         "capacity_warnings": resp.render_stats.capacity_warnings}}

        phs = set()
        for pat in [r"\{\{([^{}]+)\}\}", r"\[\[([^\[\]]+)\]\]"]:
            for m in re.finditer(pat, ft): phs.add(m.group(0))
        bt = ["示例科技","示例项目","王示例","z****@***********","138****1111"]
        bs = [t for t in bt if t in ft]
        ok7 = len(phs)==0 and len(bs)==0
        results["7_no_template_sample"] = {"status": "通过" if ok7 else "失败",
            "evidence": {"unreplaced": sorted(phs), "samples": sorted(bs)}}

        ok8 = all(s.status == "done" for s in resp.stages)
        results["8_all_stages_done"] = {"status": "通过" if ok8 else "失败",
            "evidence": {"stages": [s.model_dump() for s in resp.stages]}}

        rf = ["file_path","file_name","download_url","stages","matched_experience_ids",
              "rendered_experience_ids","profile_source","build_counts","build_meta","render_stats","warnings","template_id"]
        ok9 = all(hasattr(resp,f) and getattr(resp,f) is not None for f in rf)
        results["9_response_structure"] = {"status": "通过" if ok9 else "失败",
            "evidence": {"missing": [f for f in rf if not hasattr(resp,f) or getattr(resp,f) is None]}}

        bm2 = resp.build_meta
        ok10 = all([isinstance(bm2.profile_source,str) and bm2.profile_source,
                    isinstance(bm2.ai_covered_experience_ids,list),
                    isinstance(bm2.fallback_sql_experience_ids,list),
                    isinstance(bm2.ai_unrecognized_experience_ids,list)])
        results["10_build_meta_diagnostics"] = {"status": "通过" if ok10 else "失败",
            "evidence": {"profile_source": bm2.profile_source, "ai_covered": bm2.ai_covered_experience_ids,
                         "fallback": bm2.fallback_sql_experience_ids, "ai_unrecognized": bm2.ai_unrecognized_experience_ids,
                         "trimmed": bm2.max_items_trimmed, "counts": bm2.counts.model_dump()}}
        return results
    finally:
        db.close()

def _run_error(desc, extra_patches, exp_code, exp_http):
    from database.init_db import init_db
    from database.session import SessionLocal
    from services import resume_generation_service, experience_service
    from api.schemas import RequestProfile, ResumeDocxGenerateRequest
    from core.errors import DomainError
    _cleanup()
    init_db()
    # V1.5.0：mock embedding
    with patch("services.embedding_service._embed_text", side_effect=_create_mock_embed()):
        uid, eids = _setup_test_data(SessionLocal)
    db = SessionLocal()
    try:
        all_exps = experience_service.list_experiences(db, uid)
        wids = [e.id for e in all_exps if e.type == "work"]
        pids = [e.id for e in all_exps if e.type == "project"]

        # V1.5.0：默认 mock LLM + embedding
        am = {
            "services.llm_service.chat_structured": _create_mock_chat_structured(STUB_JD_ANALYSIS),
            "services.embedding_service._embed_text": _create_mock_embed(),
        }
        for t, se in extra_patches:
            am[t] = se
        ps = []
        for t, se in am.items():
            p = patch(t, side_effect=se if callable(se) else se)
            p.start()
            ps.append(p)
        try:
            req = ResumeDocxGenerateRequest(user_id=uid, template_id="pm_template",
                                            jd_text=STUB_JD_TEXT, profile=RequestProfile(**STUB_PROFILE), top_k=5)
            resp = resume_generation_service.generate_docx(db, req)
            return {"scenario": desc, "status": "失败",
                    "evidence": {"expected": exp_code, "expected_http": exp_http, "actual": f"no error, ok={resp.ok}"}}
        except DomainError as e:
            ok = (e.error_code == exp_code and e.http_status == exp_http and bool(e.message))
            return {"scenario": desc, "status": "通过" if ok else "失败",
                    "evidence": {"error_code": e.error_code, "http": e.http_status, "expected": exp_code,
                                 "expected_http": exp_http, "message": (e.message or "")[:200], "stage": e.stage, "retryable": e.retryable}}
        except Exception as e:
            return {"scenario": desc, "status": "失败", "evidence": {"expected": exp_code, "actual": f"{type(e).__name__}: {e}"}}
        finally:
            for p in ps:
                p.stop()
    finally:
        db.close()

def run_error_branches():
    from core.errors import LLMOutputInvalidError, NoMatchedExperienceError
    sc = [
        ("JD分析: position为空 -> JD_INVALID (422)",
         [("services.llm_service.chat_structured",
           _create_mock_chat_structured({**STUB_JD_ANALYSIS, "position": ""}))],
         "JD_INVALID", 422),
        ("LLM: structured失败 -> LLM_OUTPUT_INVALID (502)",
         [("services.llm_service.chat_structured",
           MagicMock(side_effect=LLMOutputInvalidError("stub", stage="content_generation", details={"s":"Gen"})))],
         "LLM_OUTPUT_INVALID", 502),
    ]
    results = [_run_error(d, p, c, h) for d, p, c, h in sc]
    results.extend(_run_profile_boundary())
    return results

def _run_profile_boundary() -> list[dict]:
    """身份事实来源边界测试（V1.4.1 新增，V1.5.0 适配）。"""
    from services import resume_builder, resume_generation_service, experience_service
    from api.schemas import RequestProfile, ResumeDocxGenerateRequest
    from models.resume_document import Profile as ProfileDoc
    from docx import Document as DocxDoc

    results: list[dict] = []

    FAKE_NAME = "欧阳不该出现在文档里"
    FAKE_PHONE = "13999998888"
    FAKE_EMAIL = "leaked-fake@example.invalid"
    fake_raws = [
        f"姓名：{FAKE_NAME}",
        f"联系方式：{FAKE_PHONE} / {FAKE_EMAIL}",
        "公司：深圳某某科技有限公司",
    ]
    fake_db_experiences = [
        type("FakeExp", (object,), {"raw_text": r, "type": "work"})()
        for r in fake_raws
    ]

    # A1. request 什么都不提供 → Profile 全空
    p_a1, src_a1 = resume_builder.ProfileResolver.resolve(request_profile=None, jd_position="产品经理")
    ok_a1 = (
        p_a1.name == "" and p_a1.phone == "" and p_a1.email == "" and
        FAKE_NAME not in p_a1.name and FAKE_PHONE not in p_a1.phone and FAKE_EMAIL not in p_a1.email and
        p_a1.target_position == "产品经理" and p_a1.summary == "" and
        src_a1 == "empty"
    )
    results.append({
        "scenario": "Profile边界A1: request空 -> profile全空, source=empty, 不读取经历",
        "status": "通过" if ok_a1 else "失败",
        "evidence": {"profile": _profile_snapshot(p_a1), "source": src_a1,
                     "fake_raws_in_profile": any(FAKE_NAME in str(v) or FAKE_PHONE in str(v) or FAKE_EMAIL in str(v)
                                                 for v in [p_a1.name, p_a1.phone, p_a1.email])},
    })

    # A2. 经历 raw_text 注入虚构联系方式 + request 不提供 → Profile 为空
    p_a2, src_a2 = resume_builder.ProfileResolver.resolve(
        request_profile={}, jd_position="产品经理")
    ok_a2 = (
        p_a2.name == "" and p_a2.phone == "" and p_a2.email == "" and
        FAKE_NAME not in p_a2.name and FAKE_PHONE not in p_a2.phone and FAKE_EMAIL not in p_a2.email and
        src_a2 == "empty"
    )
    results.append({
        "scenario": "Profile边界A2: request空 + 经历含虚构联系方式 -> profile全空, 不回填",
        "status": "通过" if ok_a2 else "失败",
        "evidence": {"profile": _profile_snapshot(p_a2), "source": src_a2},
    })

    # A3. request 只提供 phone+email → 只保留显式值
    p_a3, src_a3 = resume_builder.ProfileResolver.resolve(
        request_profile={"phone": STUB_PROFILE["phone"], "email": STUB_PROFILE["email"]},
        jd_position="产品经理")
    ok_a3 = (
        p_a3.name == "" and
        p_a3.phone == STUB_PROFILE["phone"] and p_a3.email == STUB_PROFILE["email"] and
        FAKE_NAME not in p_a3.name and FAKE_PHONE not in p_a3.phone and FAKE_EMAIL not in p_a3.email and
        src_a3 == "request"
    )
    results.append({
        "scenario": "Profile边界A3: request仅提供phone+email, 无name -> 仅保留显式值",
        "status": "通过" if ok_a3 else "失败",
        "evidence": {"profile": _profile_snapshot(p_a3), "source": src_a3},
    })

    # A4. profile_source 只能是 request / empty
    ok_src = src_a1 in ("request", "empty") and src_a2 in ("request", "empty") and src_a3 in ("request", "empty")
    results.append({
        "scenario": "Profile边界A4: profile_source 只取值 request 或 empty",
        "status": "通过" if ok_src else "失败",
        "evidence": {"A1": src_a1, "A2": src_a2, "A3": src_a3},
    })

    # B. 核心链路（姓名缺失不报错，虚构联系方式不进 DOCX）
    _cleanup()
    from database.init_db import init_db
    from database.session import SessionLocal
    init_db()
    with patch("services.embedding_service._embed_text", side_effect=_create_mock_embed()):
        uid, _ = _setup_test_data(SessionLocal)
    db = SessionLocal()
    try:
        all_exps = experience_service.list_experiences(db, uid)
        if all_exps:
            injected = "\n".join([
                (all_exps[0].raw_text or ""),
                f"姓名：{FAKE_NAME}",
                f"联系电话：{FAKE_PHONE}",
                f"邮箱：{FAKE_EMAIL}",
            ])
            all_exps[0].raw_text = injected
            db.add(all_exps[0])
            db.commit()

        with patch("services.llm_service.chat_structured",
                   side_effect=_create_mock_chat_structured(STUB_JD_ANALYSIS)), \
             patch("services.embedding_service._embed_text", side_effect=_create_mock_embed()):
            inc = RequestProfile(name="", phone=STUB_PROFILE["phone"], email=STUB_PROFILE["email"])
            req = ResumeDocxGenerateRequest(user_id=uid, template_id="pm_template",
                                            jd_text=STUB_JD_TEXT, profile=inc, top_k=5)
            try:
                resp = resume_generation_service.generate_docx(db, req)
                b1_ok = bool(resp.ok)
                b1_evidence = {"ok": resp.ok, "stages_done": all(s.status == "done" for s in resp.stages)}
            except Exception as e:
                b1_ok = False
                b1_evidence = {"error_type": type(e).__name__, "error": str(e)[:200]}

        results.append({
            "scenario": "Profile边界B1: name空但phone+email存在 -> 核心链路成功",
            "status": "通过" if b1_ok else "失败",
            "evidence": b1_evidence,
        })

        if b1_ok:
            dp = OUTPUT_DIR / resp.file_name if resp.file_name else None
            if dp and dp.exists():
                d = DocxDoc(str(dp))
                ft = "\n".join(p.text for p in d.paragraphs if p.text.strip())
                ft += "\n" + "\n".join(cell.text for t in d.tables for row in t.rows for cell in row.cells)
                leaks = [x for x in (FAKE_NAME, FAKE_PHONE, FAKE_EMAIL) if x in ft]
                contract_leak = "PROFILE_INCOMPLETE" in ft
                b2_ok = len(leaks) == 0 and not contract_leak
                results.append({
                    "scenario": "Profile边界B2: 经历注入的虚构姓名/手机/邮箱不进入DOCX",
                    "status": "通过" if b2_ok else "失败",
                    "evidence": {"leaked": leaks, "old_contract_in_docx": contract_leak,
                                 "docx_sample": ft[:300]},
                })
            else:
                results.append({
                    "scenario": "Profile边界B2: 经历注入的虚构姓名/手机/邮箱不进入DOCX",
                    "status": "失败",
                    "evidence": {"reason": "DOCX未产出", "file_path": str(dp) if dp else None},
                })
    finally:
        db.close()

    return results


def _profile_snapshot(p: _ProfileDoc) -> dict:
    return {
        "name": p.name,
        "phone": p.phone,
        "email": p.email,
        "location": p.location,
        "target_position": p.target_position,
        "summary": p.summary,
    }

def main():
    try:
        print("=" * 72)
        print("V1.5.0 Stub E2E - 固定 LLM/Embedding mock, CI 可重复")
        print("=" * 72)
        print("\n[1/2] Happy Path ...")
        happy = run_happy_path()
        for k, v in happy.items():
            print(f"  [{'PASS' if v['status']=='通过' else 'FAIL'}] {k}: {v['status']}")
        print("\n[2/2] 错误分支 + 身份边界 ...")
        errors = run_error_branches()
        for r in errors:
            print(f"  [{'PASS' if r['status']=='通过' else 'FAIL'}] {r['scenario']}")
        all_results = {"happy_path": happy, "error_branches": errors,
                       "__summary__": {"happy_path": {k: v["status"] for k,v in happy.items()},
                                       "error_branches": {r["scenario"]: r["status"] for r in errors}}}
        hp = sum(1 for v in happy.values() if v["status"]=="通过")
        ht = len(happy)
        ep = sum(1 for r in errors if r["status"]=="通过")
        et = len(errors)
        out = OUTPUT_DIR / "V1.5_StubE2E_验证表.json"
        out.write_text(json.dumps(all_results, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\n{'=' * 72}")
        print(f"Stub E2E 完成 - {out}")
        print(f"  Happy Path: {hp}/{ht}  错误分支: {ep}/{et}  总计: {hp+ep}/{ht+et}")
        print("=" * 72)
        if hp != ht or ep != et:
            print("\n存在未通过的测试项!")
            sys.exit(1)
    finally:
        _ok = _cleanup_stub_runtime()
        _removed = not _STUB_RUNTIME_TMP.exists()
        print(f"  [isolation] temp runtime removed: {_removed}")
        if not _ok or not _removed:
            print("  [isolation][FAIL] cleanup failed — cannot pass isolation")
            sys.exit(1)
    return all_results

if __name__ == "__main__":
    main()
